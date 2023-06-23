import json
from docx import Document
from docx2txt import process


def extract_table_data(table):
    table_data = []

    # Iterate over each row in the table
    for row in table.rows:
        # Extract the text from each cell in the row
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)

    return table_data


def extract_paragraphs(doc):
    paragraphs = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)

    return paragraphs


def detect_tables_and_paragraphs(doc_file):
    if doc_file.endswith('.docx'):
        doc = Document(doc_file)
        tables = []

        for table in doc.tables:
            # Check if the table has any rows and columns
            if table.rows and table.columns:
                table_data = extract_table_data(table)
                tables.append(table_data)

        paragraphs = extract_paragraphs(doc)

        return tables, paragraphs

    elif doc_file.endswith('.doc'):
        text = process(doc_file)
        # Perform table extraction from the extracted text
        # You can use regex or any other method to identify and extract tables from the text
        # Here, you can implement the logic to extract tables from the text
        # and return them in a similar format as for DOCX files
        tables = []
        paragraphs = []

        return tables, paragraphs

    else:
        raise ValueError("Invalid file format. Supported formats: .docx, .doc")


# Usage example
if __name__ == '__main__':
    import sys

    if len(sys.argv) < 2:
        print('Usage: python tableExtractor.py <doc_file_path>')
        sys.exit(1)

    doc_file_path = sys.argv[1]
    detected_tables, detected_paragraphs = detect_tables_and_paragraphs(doc_file_path)

    # Convert the tables to a list of dictionaries
    converted_tables = []

    for table in detected_tables:
        headers = table[0]
        data = []

        for row in table[1:]:
            row_dict = {headers[i]: row[i] for i in range(min(len(headers), len(row)))}
            data.append(row_dict)

        converted_tables.append(data)

    # Convert the tables and paragraphs to JSON and print the result
    json_result = {
        'tables': converted_tables,
        'paragraphs': detected_paragraphs
    }

    json_output = json.dumps(json_result)
    print(json_output)
